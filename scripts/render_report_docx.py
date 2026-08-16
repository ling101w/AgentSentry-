#!/usr/bin/env python3
"""Render the Chinese Markdown evaluation report as a readable DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    _configure_styles(document)

    lines = args.input.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = document.add_paragraph(style="No Spacing")
            paragraph.paragraph_format.left_indent = Cm(0.5)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            index += 1
            continue
        if _is_table_row(line) and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            rows = []
            rows.append(_split_table_row(line))
            index += 2
            while index < len(lines) and _is_table_row(lines[index]) and lines[index].strip():
                rows.append(_split_table_row(lines[index]))
                index += 1
            _add_table(document, rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            paragraph.add_run(_clean_inline(heading.group(2)))
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue
        if line.strip() == "---":
            document.add_paragraph("", style="No Spacing")
            index += 1
            continue
        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            document.add_paragraph(_clean_inline(bullet.group(1)), style="List Bullet")
            index += 1
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            document.add_paragraph(_clean_inline(numbered.group(1)), style="List Number")
            index += 1
            continue
        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run(_clean_inline(line))
        index += 1

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    return 0


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name, size in (("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)):
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True


def _clean_inline(value: str) -> str:
    value = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", value)
    value = value.replace("**", "").replace("`", "")
    return value


def _is_table_row(value: str) -> bool:
    return value.strip().startswith("|") and value.strip().endswith("|")


def _is_table_separator(value: str) -> bool:
    if not _is_table_row(value):
        return False
    cells = _split_table_row(value)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def _split_table_row(value: str) -> list[str]:
    return [cell.strip() for cell in value.strip().strip("|").split("|")]


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(columns):
            text = _clean_inline(row[column_index]) if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
        if row_index == 0:
            for cell in table.rows[row_index].cells:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    document.add_paragraph("", style="No Spacing")


if __name__ == "__main__":
    raise SystemExit(main())
